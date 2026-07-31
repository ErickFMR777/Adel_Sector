"""
estudio_sector.py — Generación del Estudio del Sector (PDF y Word).

Produce el documento con la estructura que exige la *Guía para la
Elaboración de Estudios del Sector V3* (septiembre de 2025) de la
Agencia Nacional de Contratación Pública – Colombia Compra Eficiente,
apartado 5.2:

    5.2.1  Aspectos generales del mercado
    5.2.2  Modelo de Abastecimiento Estratégico
    5.2.3  Comportamiento del gasto histórico — Estudio de la demanda
    5.2.4  Estudio de la oferta
    5.2.5  Estudio de mercado (análisis de precios)
    5.2.6  Conclusiones del Estudio del Sector

El análisis estadístico sigue el apartado 8 de la misma guía:
tendencia central (media, mediana, moda), dispersión (varianza,
desviación estándar, coeficiente de variación), posición (cuartiles) y
**manejo de datos atípicos por el criterio del rango intercuartílico
(RIC)**, con el correspondiente recálculo de estadísticas ajustadas —
que es justamente lo que la guía recomienda, porque el promedio simple
es sensible a valores extremos y deja de representar al mercado.

Importante sobre el alcance: este módulo automatiza los componentes que
**se pueden derivar de los datos de SECOP** (gasto histórico, oferentes,
precios). Los apartados que dependen del criterio de la entidad —
contexto regulatorio del bien concreto, requisitos habilitantes, riesgos
— se emiten como secciones guiadas para que la entidad las complete; no
se inventan.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from typing import Any, Optional

import pandas as pd

from config import resolver_fuente_pdf

logger = logging.getLogger(__name__)

# Texto que se inserta donde la entidad debe aportar criterio propio.
_POR_COMPLETAR = "[Por completar por la Entidad Estatal]"


# ════════════════════════════════════════════════════════════
# 1. CONTEXTO DEL ESTUDIO
# ════════════════════════════════════════════════════════════


@dataclass
class ContextoEstudio:
    """Datos de encabezado que no salen de SECOP y aporta la entidad."""

    objeto: str = ""
    entidad: str = ""
    departamento: str = ""
    municipio: str = ""
    modalidad_prevista: str = ""
    codigo_unspsc: str = ""
    elaborado_por: str = ""
    observaciones: str = ""
    # Trazabilidad de la consulta que originó el estudio.
    filtros: dict[str, Any] = field(default_factory=dict)
    fuentes: list[str] = field(default_factory=list)
    consultado_en: Optional[datetime] = None


# ════════════════════════════════════════════════════════════
# 2. ESTADÍSTICA DESCRIPTIVA (apartado 8 de la guía)
# ════════════════════════════════════════════════════════════


def calcular_estadisticas(valores: pd.Series) -> dict[str, Any]:
    """Calcula el bloque estadístico que pide la guía.

    Incluye tendencia central, dispersión, posición y la detección de
    atípicos por rango intercuartílico, más las estadísticas ajustadas
    (recalculadas sin los atípicos).

    Args:
        valores: Serie numérica con los valores de los contratos.

    Returns:
        Diccionario con todas las medidas. Las claves ``ajustadas`` y
        ``atipicos`` describen el tratamiento de valores extremos.
    """
    serie = pd.to_numeric(valores, errors="coerce").dropna()
    serie = serie[serie > 0]

    if serie.empty:
        return {"n": 0}

    q1 = float(serie.quantile(0.25))
    q2 = float(serie.quantile(0.50))
    q3 = float(serie.quantile(0.75))
    ric = q3 - q1

    # Criterio del rango intercuartílico recomendado por la guía.
    limite_inferior = q1 - 1.5 * ric
    limite_superior = q3 + 1.5 * ric

    mascara_atipicos = (serie < limite_inferior) | (serie > limite_superior)
    atipicos = serie[mascara_atipicos]
    ajustada = serie[~mascara_atipicos]

    media = float(serie.mean())
    desviacion = float(serie.std(ddof=1)) if len(serie) > 1 else 0.0

    modas = serie.mode()
    moda = float(modas.iloc[0]) if not modas.empty else None

    def _bloque(datos: pd.Series) -> dict[str, Any]:
        if datos.empty:
            return {"n": 0}
        prom = float(datos.mean())
        desv = float(datos.std(ddof=1)) if len(datos) > 1 else 0.0
        return {
            "n": int(len(datos)),
            "media": prom,
            "mediana": float(datos.median()),
            "minimo": float(datos.min()),
            "maximo": float(datos.max()),
            "desviacion": desv,
            "coef_variacion": (desv / prom * 100) if prom else 0.0,
            "total": float(datos.sum()),
        }

    return {
        "n": int(len(serie)),
        "media": media,
        "mediana": q2,
        "moda": moda,
        "minimo": float(serie.min()),
        "maximo": float(serie.max()),
        "rango": float(serie.max() - serie.min()),
        "suma": float(serie.sum()),
        "varianza": float(serie.var(ddof=1)) if len(serie) > 1 else 0.0,
        "desviacion": desviacion,
        "coef_variacion": (desviacion / media * 100) if media else 0.0,
        "q1": q1,
        "q2": q2,
        "q3": q3,
        "ric": ric,
        "p10": float(serie.quantile(0.10)),
        "p90": float(serie.quantile(0.90)),
        "limite_inferior": limite_inferior,
        "limite_superior": limite_superior,
        "atipicos": {
            "n": int(len(atipicos)),
            "pct": len(atipicos) / len(serie) * 100,
            "valores": [float(v) for v in atipicos.nlargest(5)],
        },
        "ajustadas": _bloque(ajustada),
    }


def interpretar_dispersion(coef_variacion: float) -> str:
    """Traduce el coeficiente de variación a una lectura en palabras.

    Se ofrece como apoyo a la redacción, no como un umbral normativo.
    """
    if coef_variacion < 15:
        return (
            "baja dispersión: los precios del mercado son homogéneos y la "
            "media es un buen estimador del precio de referencia"
        )
    if coef_variacion < 40:
        return (
            "dispersión moderada: conviene contrastar la media con la "
            "mediana antes de fijar el precio de referencia"
        )
    return (
        "dispersión alta: los precios son heterogéneos, por lo que la "
        "mediana y las estadísticas ajustadas representan mejor al mercado "
        "que el promedio simple"
    )


# ════════════════════════════════════════════════════════════
# 3. ANÁLISIS POR COMPONENTE
# ════════════════════════════════════════════════════════════


def _distribucion(df: pd.DataFrame, columna: str, tope: int = 10) -> pd.DataFrame:
    """Frecuencia y valor contratado agrupados por una columna."""
    if columna not in df.columns or df[columna].isna().all():
        return pd.DataFrame()

    agrupado = (
        df.groupby(df[columna].fillna("No informado").astype(str))
        .agg(contratos=("valor_del_contrato", "size"),
             valor_total=("valor_del_contrato", "sum"))
        .sort_values("valor_total", ascending=False)
    )
    total = agrupado["valor_total"].sum()
    agrupado["pct_valor"] = (
        agrupado["valor_total"] / total * 100 if total else 0
    )
    return agrupado.head(tope).reset_index().rename(columns={columna: "categoria"})


def analizar_demanda(df: pd.DataFrame) -> dict[str, Any]:
    """Componente 5.2.3 — comportamiento del gasto histórico.

    Responde a las preguntas de la guía: cómo ha adquirido el Estado este
    bien o servicio, con qué modalidades, en qué cantidades y con qué
    patrón temporal.
    """
    resultado: dict[str, Any] = {
        "total_contratos": len(df),
        "valor_total": float(
            pd.to_numeric(df["valor_del_contrato"], errors="coerce").sum()
        ),
        "modalidades": _distribucion(df, "modalidad_de_contratacion"),
        "tipos_contrato": _distribucion(df, "tipo_de_contrato"),
        "entidades": _distribucion(df, "nombre_entidad", tope=15),
        "ciudades": _distribucion(df, "ciudad", tope=10),
        "estados": _distribucion(df, "estado_contrato"),
    }

    # Comportamiento temporal: detecta estacionalidad y concentración anual.
    if "fecha_inicio" in df.columns:
        fechas = pd.to_datetime(df["fecha_inicio"], errors="coerce")
        validas = fechas.dropna()
        if not validas.empty:
            por_anio = (
                df.assign(_anio=fechas.dt.year)
                .dropna(subset=["_anio"])
                .groupby("_anio")
                .agg(contratos=("valor_del_contrato", "size"),
                     valor_total=("valor_del_contrato", "sum"))
                .reset_index()
                .rename(columns={"_anio": "anio"})
            )
            por_anio["anio"] = por_anio["anio"].astype(int)
            resultado["por_anio"] = por_anio

            por_mes = (
                df.assign(_mes=fechas.dt.month)
                .dropna(subset=["_mes"])
                .groupby("_mes")
                .agg(contratos=("valor_del_contrato", "size"))
                .reset_index()
                .rename(columns={"_mes": "mes"})
            )
            por_mes["mes"] = por_mes["mes"].astype(int)
            resultado["por_mes"] = por_mes

            resultado["periodo"] = (validas.min(), validas.max())

    return resultado


def analizar_oferta(df: pd.DataFrame) -> dict[str, Any]:
    """Componente 5.2.4 — estudio de la oferta.

    Identifica los proveedores que ya han atendido esta necesidad y mide
    la concentración del mercado, insumo para valorar si hay competencia
    suficiente y para definir requisitos habilitantes proporcionados.
    """
    resultado: dict[str, Any] = {}

    if "proveedor_adjudicado" not in df.columns:
        return {"proveedores": pd.DataFrame(), "n_proveedores": 0}

    proveedores = df["proveedor_adjudicado"].fillna("").astype(str)
    validos = df[proveedores.str.strip().ne("") & proveedores.str.lower().ne("nan")]

    resultado["n_proveedores"] = int(
        validos["proveedor_adjudicado"].nunique()
    ) if not validos.empty else 0
    resultado["proveedores"] = _distribucion(
        validos, "proveedor_adjudicado", tope=20
    )

    # Concentración: cuota del mayor proveedor y de los cinco primeros.
    tabla = resultado["proveedores"]
    if not tabla.empty:
        resultado["cuota_lider"] = float(tabla.iloc[0]["pct_valor"])
        resultado["cuota_top5"] = float(tabla.head(5)["pct_valor"].sum())
        resultado["lider"] = str(tabla.iloc[0]["categoria"])

    return resultado


def analizar_mercado(df: pd.DataFrame) -> dict[str, Any]:
    """Componente 5.2.5 — estudio de mercado (análisis de precios)."""
    estadisticas = calcular_estadisticas(df["valor_del_contrato"])
    return {
        "estadisticas": estadisticas,
        "interpretacion": (
            interpretar_dispersion(estadisticas.get("coef_variacion", 0))
            if estadisticas.get("n") else ""
        ),
    }


def construir_estudio(
    df: pd.DataFrame, contexto: ContextoEstudio
) -> dict[str, Any]:
    """Ensambla todos los componentes del Estudio del Sector."""
    df = df.copy()
    df["valor_del_contrato"] = pd.to_numeric(
        df.get("valor_del_contrato"), errors="coerce"
    )

    return {
        "contexto": contexto,
        "generado_en": datetime.now(),
        "demanda": analizar_demanda(df),
        "oferta": analizar_oferta(df),
        "mercado": analizar_mercado(df),
        "muestra": df,
    }


# ════════════════════════════════════════════════════════════
# 4. FORMATEO
# ════════════════════════════════════════════════════════════


def cop(valor: Optional[float]) -> str:
    """Formatea un valor en pesos colombianos."""
    if valor is None or pd.isna(valor):
        return "N/D"
    return f"${valor:,.0f}".replace(",", ".")


def _pct(valor: float) -> str:
    return f"{valor:,.1f}%".replace(",", ".")


_MESES = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]


def _resumen_filtros(contexto: ContextoEstudio) -> str:
    """Describe en una línea los filtros que originaron la muestra."""
    partes = [f"{k}: {v}" for k, v in contexto.filtros.items() if v]
    return "; ".join(partes) if partes else "sin filtros adicionales"


# ════════════════════════════════════════════════════════════
# 5. EXPORTACIÓN A WORD
# ════════════════════════════════════════════════════════════


def exportar_docx(estudio: dict[str, Any]) -> bytes:
    """Genera el Estudio del Sector en formato Word (.docx).

    Se emite en Word además de PDF porque el documento suele integrarse
    en los estudios previos de la entidad, donde hay que seguir editando.

    Returns:
        Contenido binario del archivo .docx.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    ctx: ContextoEstudio = estudio["contexto"]
    demanda = estudio["demanda"]
    oferta = estudio["oferta"]
    mercado = estudio["mercado"]
    est = mercado["estadisticas"]

    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(10.5)

    def titulo(texto: str, nivel: int = 1) -> None:
        h = doc.add_heading(texto, level=nivel)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0B, 0x35, 0x66)

    def parrafo(texto: str, negrita: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = negrita

    def tabla_desde_df(
        datos: pd.DataFrame, encabezados: list[str], columnas: list[str]
    ) -> None:
        if datos.empty:
            parrafo("Sin información disponible para este componente.")
            return
        t = doc.add_table(rows=1, cols=len(columnas))
        t.style = "Light Grid Accent 1"
        for i, texto in enumerate(encabezados):
            celda = t.rows[0].cells[i]
            celda.text = texto
            for p in celda.paragraphs:
                for run in p.runs:
                    run.bold = True
        for _, fila in datos.iterrows():
            celdas = t.add_row().cells
            for i, col in enumerate(columnas):
                valor = fila[col]
                if col in ("valor_total",):
                    celdas[i].text = cop(valor)
                elif col in ("pct_valor",):
                    celdas[i].text = _pct(valor)
                else:
                    celdas[i].text = str(valor)

    # ── Portada ──
    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = encabezado.add_run("ESTUDIO DEL SECTOR")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x35, 0x66)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Elaborado conforme a la Guía para la Elaboración de Estudios del "
        "Sector V3 (2025)\nAgencia Nacional de Contratación Pública — "
        "Colombia Compra Eficiente"
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_paragraph()

    ficha = doc.add_table(rows=0, cols=2)
    ficha.style = "Light List Accent 1"
    for etiqueta, valor in (
        ("Objeto a contratar", ctx.objeto or _POR_COMPLETAR),
        ("Entidad Estatal", ctx.entidad or _POR_COMPLETAR),
        ("Departamento", ctx.departamento or "Nacional"),
        ("Municipio", ctx.municipio or "—"),
        ("Modalidad prevista", ctx.modalidad_prevista or _POR_COMPLETAR),
        ("Código UNSPSC", ctx.codigo_unspsc or _POR_COMPLETAR),
        ("Elaborado por", ctx.elaborado_por or _POR_COMPLETAR),
        ("Fecha de elaboración", f"{estudio['generado_en']:%d/%m/%Y %H:%M}"),
    ):
        fila = ficha.add_row().cells
        fila[0].text = etiqueta
        for p in fila[0].paragraphs:
            for r in p.runs:
                r.bold = True
        fila[1].text = str(valor)

    doc.add_page_break()

    # ── 1. Objetivo y alcance ──
    titulo("1. Objetivo y alcance", 1)
    parrafo(
        "El presente documento contiene el análisis del sector económico "
        "relativo al objeto del Proceso de Contratación, en cumplimiento "
        "del artículo 2.2.1.1.1.6.1 del Decreto 1082 de 2015, que obliga a "
        "la Entidad Estatal a analizar el sector desde las perspectivas "
        "legal, comercial, financiera, organizacional, técnica y de "
        "análisis de riesgo."
    )
    parrafo(
        f"El estudio se construye sobre {demanda['total_contratos']:,} "
        "contratos de referencia extraídos del Sistema Electrónico de "
        "Contratación Pública (SECOP), que representan "
        f"{cop(demanda['valor_total'])} en contratación pública."
        .replace(",", ".")
    )

    # ── 2. Aspectos generales del mercado ──
    titulo("2. Aspectos generales del mercado (5.2.1)", 1)
    parrafo(
        "La guía recomienda analizar el sector desde los contextos "
        "económico, técnico, regulatorio y de mercado. Los componentes "
        "cuantitativos se desarrollan en los numerales siguientes a partir "
        "de la contratación histórica registrada en SECOP."
    )
    titulo("2.1 Contexto económico", 2)
    parrafo(ctx.observaciones or _POR_COMPLETAR)
    titulo("2.2 Contexto técnico", 2)
    parrafo(_POR_COMPLETAR)
    titulo("2.3 Contexto regulatorio", 2)
    parrafo(_POR_COMPLETAR)

    # ── 3. Demanda ──
    titulo("3. Comportamiento del gasto histórico — Estudio de la demanda (5.2.3)", 1)
    parrafo(
        "Analiza cómo han adquirido las Entidades Estatales este bien, obra "
        "o servicio: modalidades empleadas, valores ejecutados, número de "
        "contratos y comportamiento temporal."
    )

    if demanda.get("periodo"):
        desde, hasta = demanda["periodo"]
        parrafo(
            f"Periodo analizado: {desde:%d/%m/%Y} a {hasta:%d/%m/%Y}.",
            negrita=True,
        )

    titulo("3.1 Modalidades de selección utilizadas", 2)
    tabla_desde_df(
        demanda["modalidades"],
        ["Modalidad", "Contratos", "Valor total", "% del valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
    )

    titulo("3.2 Tipos de contrato", 2)
    tabla_desde_df(
        demanda["tipos_contrato"],
        ["Tipo de contrato", "Contratos", "Valor total", "% del valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
    )

    titulo("3.3 Entidades contratantes", 2)
    tabla_desde_df(
        demanda["entidades"],
        ["Entidad", "Contratos", "Valor total", "% del valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
    )

    if "por_anio" in demanda:
        titulo("3.4 Comportamiento anual", 2)
        tabla_desde_df(
            demanda["por_anio"],
            ["Año", "Contratos", "Valor total"],
            ["anio", "contratos", "valor_total"],
        )

    if "por_mes" in demanda:
        titulo("3.5 Estacionalidad", 2)
        mes_df = demanda["por_mes"].copy()
        mes_df["nombre"] = mes_df["mes"].map(lambda m: _MESES[m])
        pico = mes_df.loc[mes_df["contratos"].idxmax()]
        parrafo(
            f"El mes con mayor número de contrataciones es "
            f"{pico['nombre']} ({int(pico['contratos'])} contratos), lo que "
            "sugiere concentrar la planeación del proceso con anterioridad "
            "a ese periodo."
        )
        tabla_desde_df(
            mes_df, ["Mes", "Contratos"], ["nombre", "contratos"]
        )

    # ── 4. Oferta ──
    titulo("4. Estudio de la oferta (5.2.4)", 1)
    parrafo(
        f"Se identificaron {oferta.get('n_proveedores', 0):,} proveedores "
        "que han atendido esta necesidad ante el Estado. Este listado es "
        "insumo para dimensionar la competencia disponible y definir "
        "requisitos habilitantes proporcionados."
        .replace(",", ".")
    )

    if oferta.get("cuota_lider") is not None:
        parrafo(
            f"Concentración del mercado: el mayor proveedor concentra el "
            f"{_pct(oferta['cuota_lider'])} del valor contratado y los cinco "
            f"primeros el {_pct(oferta['cuota_top5'])}.",
            negrita=True,
        )

    titulo("4.1 Proveedores identificados", 2)
    tabla_desde_df(
        oferta.get("proveedores", pd.DataFrame()),
        ["Proveedor", "Contratos", "Valor total", "% del valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
    )

    # ── 5. Mercado / precios ──
    titulo("5. Estudio de mercado — análisis de precios (5.2.5)", 1)

    if not est.get("n"):
        parrafo("No hay valores de contrato suficientes para el análisis.")
    else:
        titulo("5.1 Estadísticas descriptivas", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for etiqueta, valor in (
            ("Número de contratos analizados", f"{est['n']:,}".replace(",", ".")),
            ("Valor total contratado", cop(est["suma"])),
            ("Media (promedio)", cop(est["media"])),
            ("Mediana", cop(est["mediana"])),
            ("Moda", cop(est["moda"])),
            ("Valor mínimo", cop(est["minimo"])),
            ("Valor máximo", cop(est["maximo"])),
            ("Rango", cop(est["rango"])),
        ):
            fila = t.add_row().cells
            fila[0].text = etiqueta
            fila[1].text = valor

        titulo("5.2 Medidas de dispersión", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for etiqueta, valor in (
            ("Desviación estándar", cop(est["desviacion"])),
            ("Varianza", f"{est['varianza']:,.0f}".replace(",", ".")),
            ("Coeficiente de variación", _pct(est["coef_variacion"])),
        ):
            fila = t.add_row().cells
            fila[0].text = etiqueta
            fila[1].text = valor
        parrafo(f"Lectura: {mercado['interpretacion']}.")

        titulo("5.3 Medidas de posición", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for etiqueta, valor in (
            ("Percentil 10", cop(est["p10"])),
            ("Primer cuartil (Q1)", cop(est["q1"])),
            ("Segundo cuartil (Q2 — mediana)", cop(est["q2"])),
            ("Tercer cuartil (Q3)", cop(est["q3"])),
            ("Percentil 90", cop(est["p90"])),
            ("Rango intercuartílico (RIC)", cop(est["ric"])),
        ):
            fila = t.add_row().cells
            fila[0].text = etiqueta
            fila[1].text = valor

        titulo("5.4 Identificación de datos atípicos", 2)
        parrafo(
            "Siguiendo el criterio del rango intercuartílico recomendado por "
            "la guía, se consideran atípicos los valores fuera del intervalo "
            f"[{cop(est['limite_inferior'])} , {cop(est['limite_superior'])}]."
        )
        parrafo(
            f"Se identificaron {est['atipicos']['n']:,} contratos atípicos "
            f"({_pct(est['atipicos']['pct'])} del total)."
            .replace(",", ".")
        )

        titulo("5.5 Estadísticas descriptivas ajustadas", 2)
        ajustadas = est["ajustadas"]
        if ajustadas.get("n"):
            parrafo(
                "Recalculadas tras excluir los valores atípicos, de acuerdo "
                "con la recomendación de la guía:"
            )
            t = doc.add_table(rows=0, cols=2)
            t.style = "Light Grid Accent 1"
            for etiqueta, valor in (
                ("Contratos considerados", f"{ajustadas['n']:,}".replace(",", ".")),
                ("Media ajustada", cop(ajustadas["media"])),
                ("Mediana ajustada", cop(ajustadas["mediana"])),
                ("Desviación estándar ajustada", cop(ajustadas["desviacion"])),
                ("Coeficiente de variación ajustado",
                 _pct(ajustadas["coef_variacion"])),
                ("Rango ajustado",
                 f"{cop(ajustadas['minimo'])} — {cop(ajustadas['maximo'])}"),
            ):
                fila = t.add_row().cells
                fila[0].text = etiqueta
                fila[1].text = valor

    # ── 6. Conclusiones ──
    titulo("6. Conclusiones del Estudio del Sector (5.2.6)", 1)

    if est.get("n"):
        ajustadas = est["ajustadas"]
        referencia = ajustadas.get("media") or est["media"]
        parrafo(
            f"a) Precio promedio del mercado: {cop(est['media'])}. "
            f"Excluidos los valores atípicos, el promedio ajustado es "
            f"{cop(ajustadas.get('media'))} y la mediana "
            f"{cop(est['mediana'])}."
        )
        parrafo(
            f"b) Precio de referencia sugerido: {cop(referencia)}. Se toma "
            "el promedio ajustado por ser el estimador menos sensible a "
            "valores extremos, conforme al apartado 8.3.5 de la guía."
        )
        parrafo(
            f"c) Rango de precios observado en el mercado: entre "
            f"{cop(est['q1'])} y {cop(est['q3'])} (recorrido intercuartílico, "
            "que agrupa el 50 % central de los contratos)."
        )
        parrafo(
            f"d) Potenciales oferentes: se identificaron "
            f"{oferta.get('n_proveedores', 0):,} proveedores con experiencia "
            "acreditable en objetos similares ante el Estado."
            .replace(",", ".")
        )
        if not demanda["modalidades"].empty:
            top_mod = demanda["modalidades"].iloc[0]
            parrafo(
                f"e) Modalidad de selección predominante: "
                f"{top_mod['categoria']}, empleada en "
                f"{int(top_mod['contratos'])} de los contratos analizados "
                f"({_pct(top_mod['pct_valor'])} del valor)."
            )

    parrafo(
        "f) Definición del presupuesto oficial: " + _POR_COMPLETAR,
    )
    parrafo(
        "g) Requisitos habilitantes y criterios diferenciales: "
        + _POR_COMPLETAR
    )
    parrafo("h) Análisis de riesgos: " + _POR_COMPLETAR)

    # ── Ficha técnica ──
    doc.add_page_break()
    titulo("Ficha técnica de la consulta", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    for etiqueta, valor in (
        ("Fuente de los datos", ", ".join(ctx.fuentes) or "SECOP"),
        ("Filtros aplicados", _resumen_filtros(ctx)),
        ("Contratos en la muestra", f"{demanda['total_contratos']:,}".replace(",", ".")),
        ("Fecha de consulta a SECOP",
         f"{ctx.consultado_en:%d/%m/%Y %H:%M}" if ctx.consultado_en else "N/D"),
        ("Fecha de generación", f"{estudio['generado_en']:%d/%m/%Y %H:%M}"),
        ("Metodología estadística",
         "Guía para la Elaboración de Estudios del Sector V3 (2025), "
         "apartado 8: atípicos por rango intercuartílico y estadísticas "
         "descriptivas ajustadas."),
    ):
        fila = t.add_row().cells
        fila[0].text = etiqueta
        for p in fila[0].paragraphs:
            for r in p.runs:
                r.bold = True
        fila[1].text = str(valor)

    nota = doc.add_paragraph()
    run = nota.add_run(
        "\nLos apartados marcados como «Por completar por la Entidad "
        "Estatal» requieren el criterio de la entidad y no pueden derivarse "
        "automáticamente de los datos de SECOP."
    )
    run.italic = True
    run.font.size = Pt(8.5)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ════════════════════════════════════════════════════════════
# 6. EXPORTACIÓN A PDF
# ════════════════════════════════════════════════════════════


def hay_soporte_pdf() -> bool:
    """Indica si el entorno puede generar el estudio en PDF."""
    try:
        import fpdf  # noqa: F401
    except ImportError:
        return False
    return resolver_fuente_pdf() is not None


def exportar_pdf(estudio: dict[str, Any]) -> bytes:
    """Genera el Estudio del Sector en PDF.

    Raises:
        RuntimeError: Si no hay ninguna fuente TrueType disponible.
    """
    from fpdf import FPDF

    fuentes = resolver_fuente_pdf()
    if fuentes is None:
        raise RuntimeError(
            "No se encontró una fuente TrueType para generar el PDF. "
            "Instala 'fonts-dejavu-core' o define PDF_FONT_DIR."
        )
    ruta_regular, ruta_negrita = fuentes

    ctx: ContextoEstudio = estudio["contexto"]
    demanda = estudio["demanda"]
    oferta = estudio["oferta"]
    mercado = estudio["mercado"]
    est = mercado["estadisticas"]

    FUENTE = "EstudioSans"
    AZUL = (11, 53, 102)
    GRIS = (90, 90, 95)
    GRIS_FONDO = (240, 242, 246)
    MARGEN = 16
    ANCHO = 210 - 2 * MARGEN

    class DocumentoPDF(FPDF):
        def header(self) -> None:
            if self.page_no() == 1:
                return
            self.set_font(FUENTE, "", 7.5)
            self.set_text_color(*GRIS)
            self.cell(0, 6, "Estudio del Sector — SECOP", align="L")
            self.set_x(-60)
            self.cell(0, 6, f"{estudio['generado_en']:%d/%m/%Y}", align="R")
            self.ln(8)

        def footer(self) -> None:
            self.set_y(-14)
            self.set_draw_color(200, 205, 212)
            self.line(MARGEN, self.get_y(), 210 - MARGEN, self.get_y())
            self.ln(1)
            self.set_font(FUENTE, "", 7)
            self.set_text_color(*GRIS)
            self.cell(
                0, 6,
                "Elaborado conforme a la Guía para la Elaboración de "
                "Estudios del Sector V3 (2025) — ANCP-CCE",
                align="L",
            )
            self.set_x(-40)
            self.cell(0, 6, f"Pág. {self.page_no()}", align="R")

    pdf = DocumentoPDF(orientation="P", unit="mm", format="A4")
    pdf.add_font(FUENTE, "", ruta_regular)
    pdf.add_font(FUENTE, "B", ruta_negrita)
    pdf.add_font(FUENTE, "I", ruta_regular)
    pdf.set_margins(MARGEN, MARGEN, MARGEN)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    def h1(texto: str) -> None:
        pdf.ln(3)
        pdf.set_fill_color(*AZUL)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(FUENTE, "B", 11)
        pdf.cell(0, 8, f"  {texto}", new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.ln(2)
        pdf.set_text_color(0, 0, 0)

    def h2(texto: str) -> None:
        pdf.ln(1.5)
        pdf.set_font(FUENTE, "B", 9.5)
        pdf.set_text_color(*AZUL)
        pdf.cell(0, 6, texto, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)

    def texto(contenido: str, tam: float = 9) -> None:
        pdf.set_font(FUENTE, "", tam)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(ANCHO, 4.6, contenido, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    def ficha(filas: list[tuple[str, str]], ancho_etiqueta: float = 62) -> None:
        pdf.set_font(FUENTE, "", 8.5)
        for i, (etiqueta, valor) in enumerate(filas):
            pdf.set_fill_color(*(GRIS_FONDO if i % 2 == 0 else (255, 255, 255)))
            y0 = pdf.get_y()
            pdf.set_font(FUENTE, "B", 8.5)
            pdf.cell(ancho_etiqueta, 6, f" {etiqueta}", fill=True)
            pdf.set_font(FUENTE, "", 8.5)
            pdf.multi_cell(
                ANCHO - ancho_etiqueta, 6, f" {valor}",
                new_x="LMARGIN", new_y="NEXT", fill=True,
            )
            if pdf.get_y() < y0:
                pdf.set_y(y0 + 6)
        pdf.ln(2)

    def tabla(
        datos: pd.DataFrame,
        encabezados: list[str],
        columnas: list[str],
        anchos: list[float],
    ) -> None:
        if datos.empty:
            texto("Sin información disponible para este componente.", 8.5)
            return

        pdf.set_font(FUENTE, "B", 8)
        pdf.set_fill_color(*AZUL)
        pdf.set_text_color(255, 255, 255)
        for encabezado, ancho in zip(encabezados, anchos):
            pdf.cell(ancho, 6.5, f" {encabezado}", fill=True, align="L")
        pdf.ln(6.5)

        pdf.set_text_color(25, 25, 25)
        for i, (_, fila) in enumerate(datos.iterrows()):
            if pdf.get_y() > 255:
                pdf.add_page()
            pdf.set_fill_color(*(GRIS_FONDO if i % 2 == 0 else (255, 255, 255)))
            pdf.set_font(FUENTE, "", 7.5)
            for col, ancho in zip(columnas, anchos):
                valor = fila[col]
                if col == "valor_total":
                    txt, alineacion = cop(valor), "R"
                elif col == "pct_valor":
                    txt, alineacion = _pct(valor), "R"
                elif col == "contratos":
                    txt, alineacion = f"{int(valor):,}".replace(",", "."), "R"
                else:
                    limite = int(ancho / 1.7)
                    txt = str(valor)
                    txt = txt[:limite] + "…" if len(txt) > limite else txt
                    alineacion = "L"
                pdf.cell(ancho, 5.6, f" {txt}", fill=True, align=alineacion)
            pdf.ln(5.6)
        pdf.ln(2)

    # ── Portada ──
    pdf.set_fill_color(*AZUL)
    pdf.rect(0, 0, 210, 52, "F")
    pdf.set_y(16)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(FUENTE, "B", 24)
    pdf.cell(0, 12, "ESTUDIO DEL SECTOR", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font(FUENTE, "", 8.5)
    pdf.cell(
        0, 5,
        "Guía para la Elaboración de Estudios del Sector V3 (2025)",
        new_x="LMARGIN", new_y="NEXT", align="C",
    )
    pdf.cell(
        0, 5,
        "Agencia Nacional de Contratación Pública — Colombia Compra Eficiente",
        new_x="LMARGIN", new_y="NEXT", align="C",
    )
    pdf.set_y(60)
    pdf.set_text_color(0, 0, 0)

    ficha([
        ("Objeto a contratar", ctx.objeto or _POR_COMPLETAR),
        ("Entidad Estatal", ctx.entidad or _POR_COMPLETAR),
        ("Departamento", ctx.departamento or "Nacional"),
        ("Municipio", ctx.municipio or "—"),
        ("Modalidad prevista", ctx.modalidad_prevista or _POR_COMPLETAR),
        ("Código UNSPSC", ctx.codigo_unspsc or _POR_COMPLETAR),
        ("Elaborado por", ctx.elaborado_por or _POR_COMPLETAR),
        ("Fecha de elaboración", f"{estudio['generado_en']:%d/%m/%Y %H:%M}"),
    ])

    # ── 1. Objetivo ──
    h1("1. Objetivo y alcance")
    texto(
        "El presente documento contiene el análisis del sector económico "
        "relativo al objeto del Proceso de Contratación, en cumplimiento del "
        "artículo 2.2.1.1.1.6.1 del Decreto 1082 de 2015, que obliga a la "
        "Entidad Estatal a analizar el sector desde las perspectivas legal, "
        "comercial, financiera, organizacional, técnica y de análisis de riesgo."
    )
    texto(
        f"El estudio se construye sobre "
        f"{demanda['total_contratos']:,}".replace(",", ".")
        + " contratos de referencia extraídos del SECOP, que representan "
        f"{cop(demanda['valor_total'])} en contratación pública."
    )

    # ── 2. Aspectos generales ──
    h1("2. Aspectos generales del mercado (5.2.1)")
    texto(
        "La guía recomienda analizar el sector desde los contextos económico, "
        "técnico, regulatorio y de mercado. Los componentes cuantitativos se "
        "desarrollan en los numerales siguientes a partir de la contratación "
        "histórica registrada en SECOP."
    )
    h2("2.1 Contexto económico")
    texto(ctx.observaciones or _POR_COMPLETAR)
    h2("2.2 Contexto técnico")
    texto(_POR_COMPLETAR)
    h2("2.3 Contexto regulatorio")
    texto(_POR_COMPLETAR)

    # ── 3. Demanda ──
    pdf.add_page()
    h1("3. Comportamiento del gasto histórico — Estudio de la demanda (5.2.3)")
    if demanda.get("periodo"):
        desde, hasta = demanda["periodo"]
        texto(f"Periodo analizado: {desde:%d/%m/%Y} a {hasta:%d/%m/%Y}.")

    h2("3.1 Modalidades de selección utilizadas")
    tabla(
        demanda["modalidades"],
        ["Modalidad", "Contratos", "Valor total", "% valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
        [78, 22, 50, 28],
    )

    h2("3.2 Tipos de contrato")
    tabla(
        demanda["tipos_contrato"],
        ["Tipo de contrato", "Contratos", "Valor total", "% valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
        [78, 22, 50, 28],
    )

    h2("3.3 Entidades contratantes")
    tabla(
        demanda["entidades"],
        ["Entidad", "Contratos", "Valor total", "% valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
        [78, 22, 50, 28],
    )

    if "por_anio" in demanda:
        h2("3.4 Comportamiento anual")
        tabla(
            demanda["por_anio"],
            ["Año", "Contratos", "Valor total"],
            ["anio", "contratos", "valor_total"],
            [40, 40, 98],
        )

    if "por_mes" in demanda:
        mes_df = demanda["por_mes"].copy()
        mes_df["nombre"] = mes_df["mes"].map(lambda m: _MESES[m])
        pico = mes_df.loc[mes_df["contratos"].idxmax()]
        h2("3.5 Estacionalidad")
        texto(
            f"El mes con mayor número de contrataciones es {pico['nombre']} "
            f"({int(pico['contratos'])} contratos), lo que sugiere concentrar "
            "la planeación del proceso con anterioridad a ese periodo."
        )

    # ── 4. Oferta ──
    pdf.add_page()
    h1("4. Estudio de la oferta (5.2.4)")
    texto(
        f"Se identificaron {oferta.get('n_proveedores', 0):,}".replace(",", ".")
        + " proveedores que han atendido esta necesidad ante el Estado. Este "
        "listado es insumo para dimensionar la competencia disponible y "
        "definir requisitos habilitantes proporcionados."
    )
    if oferta.get("cuota_lider") is not None:
        texto(
            f"Concentración del mercado: el mayor proveedor concentra el "
            f"{_pct(oferta['cuota_lider'])} del valor contratado y los cinco "
            f"primeros el {_pct(oferta['cuota_top5'])}."
        )
    h2("4.1 Proveedores identificados")
    tabla(
        oferta.get("proveedores", pd.DataFrame()),
        ["Proveedor", "Contratos", "Valor total", "% valor"],
        ["categoria", "contratos", "valor_total", "pct_valor"],
        [78, 22, 50, 28],
    )

    # ── 5. Mercado ──
    pdf.add_page()
    h1("5. Estudio de mercado — análisis de precios (5.2.5)")

    if not est.get("n"):
        texto("No hay valores de contrato suficientes para el análisis.")
    else:
        h2("5.1 Estadísticas descriptivas")
        ficha([
            ("Número de contratos analizados", f"{est['n']:,}".replace(",", ".")),
            ("Valor total contratado", cop(est["suma"])),
            ("Media (promedio)", cop(est["media"])),
            ("Mediana", cop(est["mediana"])),
            ("Moda", cop(est["moda"])),
            ("Valor mínimo", cop(est["minimo"])),
            ("Valor máximo", cop(est["maximo"])),
        ])

        h2("5.2 Medidas de dispersión")
        ficha([
            ("Desviación estándar", cop(est["desviacion"])),
            ("Coeficiente de variación", _pct(est["coef_variacion"])),
        ])
        texto(f"Lectura: {mercado['interpretacion']}.", 8.5)

        h2("5.3 Medidas de posición")
        ficha([
            ("Percentil 10", cop(est["p10"])),
            ("Primer cuartil (Q1)", cop(est["q1"])),
            ("Mediana (Q2)", cop(est["q2"])),
            ("Tercer cuartil (Q3)", cop(est["q3"])),
            ("Percentil 90", cop(est["p90"])),
            ("Rango intercuartílico (RIC)", cop(est["ric"])),
        ])

        h2("5.4 Identificación de datos atípicos")
        texto(
            "Siguiendo el criterio del rango intercuartílico recomendado por "
            "la guía, se consideran atípicos los valores fuera del intervalo "
            f"[{cop(est['limite_inferior'])} , {cop(est['limite_superior'])}]. "
            f"Se identificaron {est['atipicos']['n']} contratos atípicos "
            f"({_pct(est['atipicos']['pct'])} del total)."
        )

        ajustadas = est["ajustadas"]
        if ajustadas.get("n"):
            h2("5.5 Estadísticas descriptivas ajustadas")
            texto(
                "Recalculadas tras excluir los valores atípicos, conforme al "
                "apartado 8.3.5 de la guía:", 8.5,
            )
            ficha([
                ("Contratos considerados", f"{ajustadas['n']:,}".replace(",", ".")),
                ("Media ajustada", cop(ajustadas["media"])),
                ("Mediana ajustada", cop(ajustadas["mediana"])),
                ("Desviación estándar ajustada", cop(ajustadas["desviacion"])),
                ("Coeficiente de variación ajustado",
                 _pct(ajustadas["coef_variacion"])),
                ("Rango ajustado",
                 f"{cop(ajustadas['minimo'])} — {cop(ajustadas['maximo'])}"),
            ])

    # ── 6. Conclusiones ──
    pdf.add_page()
    h1("6. Conclusiones del Estudio del Sector (5.2.6)")

    if est.get("n"):
        ajustadas = est["ajustadas"]
        referencia = ajustadas.get("media") or est["media"]
        texto(
            f"a) Precio promedio del mercado: {cop(est['media'])}. Excluidos "
            f"los valores atípicos, el promedio ajustado es "
            f"{cop(ajustadas.get('media'))} y la mediana {cop(est['mediana'])}."
        )
        texto(
            f"b) Precio de referencia sugerido: {cop(referencia)}. Se toma el "
            "promedio ajustado por ser el estimador menos sensible a valores "
            "extremos, conforme al apartado 8.3.5 de la guía."
        )
        texto(
            f"c) Rango de precios observado en el mercado: entre "
            f"{cop(est['q1'])} y {cop(est['q3'])} (recorrido intercuartílico, "
            "que agrupa el 50 % central de los contratos)."
        )
        texto(
            "d) Potenciales oferentes: se identificaron "
            f"{oferta.get('n_proveedores', 0):,}".replace(",", ".")
            + " proveedores con experiencia acreditable en objetos similares "
            "ante el Estado."
        )
        if not demanda["modalidades"].empty:
            top_mod = demanda["modalidades"].iloc[0]
            texto(
                f"e) Modalidad de selección predominante: "
                f"{top_mod['categoria']}, empleada en "
                f"{int(top_mod['contratos'])} de los contratos analizados "
                f"({_pct(top_mod['pct_valor'])} del valor)."
            )

    texto("f) Definición del presupuesto oficial: " + _POR_COMPLETAR)
    texto("g) Requisitos habilitantes y criterios diferenciales: " + _POR_COMPLETAR)
    texto("h) Análisis de riesgos: " + _POR_COMPLETAR)

    # ── Ficha técnica ──
    h1("Ficha técnica de la consulta")
    ficha([
        ("Fuente de los datos", ", ".join(ctx.fuentes) or "SECOP"),
        ("Filtros aplicados", _resumen_filtros(ctx)),
        ("Contratos en la muestra",
         f"{demanda['total_contratos']:,}".replace(",", ".")),
        ("Fecha de consulta a SECOP",
         f"{ctx.consultado_en:%d/%m/%Y %H:%M}" if ctx.consultado_en else "N/D"),
        ("Fecha de generación", f"{estudio['generado_en']:%d/%m/%Y %H:%M}"),
    ], ancho_etiqueta=52)

    pdf.set_font(FUENTE, "I", 7.5)
    pdf.set_text_color(*GRIS)
    pdf.multi_cell(
        ANCHO, 4,
        "Los apartados marcados como «Por completar por la Entidad Estatal» "
        "requieren el criterio de la entidad y no pueden derivarse "
        "automáticamente de los datos de SECOP.",
        new_x="LMARGIN", new_y="NEXT",
    )

    return bytes(pdf.output())
